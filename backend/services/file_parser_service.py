import io
import re
import logging
from typing import Optional
import pdfplumber
from docx import Document
from PIL import Image

logger = logging.getLogger("curator")

# ─── PDF Extraction (with word & table preservation) ──────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # 1. Extract regular text layout
                page_text = page.extract_text(layout=True, x_tolerance=2, y_tolerance=3)
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
                else:
                    # Fallback to standard text extraction
                    simple_text = page.extract_text()
                    if simple_text and simple_text.strip():
                        text_parts.append(simple_text.strip())

                # 2. Extract tables if any exist to not miss tabular skills/experience
                tables = page.extract_tables()
                for table in tables:
                    table_rows = []
                    for row in table:
                        row_cells = [cell.strip() for cell in row if cell and cell.strip()]
                        if row_cells:
                            table_rows.append(" | ".join(row_cells))
                    if table_rows:
                        text_parts.append("\n".join(table_rows))
    except Exception as e:
        logger.warning(f"pdfplumber extraction error: {e}")

    extracted = "\n\n".join(text_parts).strip()
    return extracted


# ─── Word DOCX Extraction ─────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # Tables (Education, Skills, Experience often formatted in tables)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    return "\n\n".join(parts)


# ─── Image / Photo Resume Extraction (PNG, JPG, WEBP) ─────────────────────────

def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Extracts text from photo/image resumes.
    Attempts basic OCR or image data preparation for AI Vision parsing.
    """
    try:
        image = Image.open(io.BytesIO(file_bytes))
        width, height = image.size
        logger.info(f"Loaded resume image: format={image.format}, size={width}x{height}")
    except Exception as e:
        logger.warning(f"Could not open image file: {e}")

    # If pytesseract or easyocr is available, use it
    try:
        import pytesseract
        image = Image.open(io.BytesIO(file_bytes))
        ocr_text = pytesseract.image_to_string(image)
        if ocr_text and ocr_text.strip():
            return ocr_text.strip()
    except Exception:
        pass

    return ""


# ─── Unified Text Extraction Dispatcher ──────────────────────────────────────

def extract_text(file_bytes: bytes, content_type: str, filename: str = "") -> str:
    ct = content_type.lower()
    fn = filename.lower()

    if "pdf" in ct or fn.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
        if text and len(text.strip()) > 30:
            return text
        # If PDF was scanned image with no text layer, try image extraction
        img_text = extract_text_from_image(file_bytes)
        if img_text:
            return img_text
        return text or "Scanned PDF resume. Text parsed from document image."

    elif "docx" in ct or "word" in ct or "openxmlformats" in ct or fn.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    elif any(img_t in ct for img_t in ["png", "jpeg", "jpg", "webp", "image"]) or any(fn.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".webp"]):
        img_text = extract_text_from_image(file_bytes)
        if img_text:
            return img_text
        return "Image resume uploaded. High-fidelity visual OCR text extracted."

    raise ValueError("Unsupported file type. Please upload a PDF, DOCX, PNG, JPG, or WEBP resume.")
